import os
import csv
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import io
from typing import Dict, Any
import tempfile
import shutil
import uuid

app = FastAPI(title="온실가스 배출량 보고서 생성 API")

# CORS 설정
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 실제 환경에서는 구체적인 도메인으로 제한
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class ReportRequest(BaseModel):
    company_name: str
    selected_report_year: str
    base_year: str
    report_sales: str
    report_employees: str
    report_sales_last_year: str
    report_employees_last_year: str
    word_table1_csv: str
    word_table2_csv: str
    word_table3_csv: str
    word_table4_csv: str
    word_table5_csv: str
    word_table6_csv: str
    word_chart1_csv: str

def csv_string_to_dataframe(csv_string: str) -> pd.DataFrame:
    """CSV 문자열을 pandas DataFrame으로 변환합니다."""
    if not csv_string.strip():
        return pd.DataFrame()
    
    try:
        # pandas의 StringIO를 사용하여 더 안전하게 파싱
        from io import StringIO
        
        # CSV 문자열을 pandas로 직접 읽기
        df = pd.read_csv(StringIO(csv_string.strip()))
        
        return df
        
    except Exception as e:
        print(f"CSV 파싱 오류: {str(e)}")
        print(f"CSV 내용 (처음 500자): {csv_string[:500]}")
        
        # 수동 파싱으로 폴백
        try:
            lines = csv_string.strip().split('\n')
            reader = csv.reader(lines)
            data = list(reader)
            
            if not data:
                return pd.DataFrame()
            
            # 첫 번째 행을 헤더로 사용
            headers = data[0]
            rows = data[1:]
            
            # 각 행의 컬럼 수를 헤더 컬럼 수에 맞춰 조정
            normalized_rows = []
            for row in rows:
                if len(row) < len(headers):
                    # 부족한 컬럼은 빈 문자열로 채움
                    row.extend([''] * (len(headers) - len(row)))
                elif len(row) > len(headers):
                    # 초과하는 컬럼은 제거
                    row = row[:len(headers)]
                normalized_rows.append(row)
            
            return pd.DataFrame(normalized_rows, columns=headers)
            
        except Exception as e2:
            print(f"수동 CSV 파싱도 실패: {str(e2)}")
            return pd.DataFrame()

def safe_float_convert(value):
    """문자열 값을 안전하게 float로 변환합니다."""
    if isinstance(value, (int, float)):
        return float(value)
    
    if isinstance(value, str):
        # 쉼표 제거 및 따옴표 제거
        cleaned_value = value.replace(',', '').replace('"', '').replace("'", "").strip()
        
        # 빈 문자열이나 0.000인 경우
        if not cleaned_value or cleaned_value == '0.000':
            return 0.0
            
        try:
            return float(cleaned_value)
        except ValueError:
            return 0.0
    
    return 0.0

def calculate_variables(request_data: ReportRequest) -> Dict[str, Any]:
    """요청 데이터를 바탕으로 필요한 변수들을 계산합니다."""
    
    try:
        # 기본 변수
        company_name = request_data.company_name
        report_year = int(request_data.selected_report_year)
        base_year = int(request_data.base_year)
        previous_year = report_year - 1
        
        # Table3 데이터에서 사업장 수 계산
        print("Table3 데이터 파싱 중...")
        table3_df = csv_string_to_dataframe(request_data.word_table3_csv)
        workplace_num = len(table3_df) if not table3_df.empty else 0
        print(f"사업장 수: {workplace_num}")
        
        # Table1 데이터에서 배출량 정보 추출
        print("Table1 데이터 파싱 중...")
        table1_df = csv_string_to_dataframe(request_data.word_table1_csv)
        print(f"Table1 DataFrame shape: {table1_df.shape}")
        print(f"Table1 columns: {table1_df.columns.tolist()}")
        
        if table1_df.empty:
            raise ValueError("Table1 데이터가 비어있습니다.")
        
        # 총합 행 찾기
        total_row = table1_df[table1_df['구분'] == '총합']
        if total_row.empty:
            raise ValueError("총합 행을 찾을 수 없습니다.")
        
        total_row = total_row.iloc[0]
        
        # 배출량 데이터 추출 및 변환
        total_emission_report_year = safe_float_convert(total_row['보고대상연도 배출량(tCO2eq)'])
        total_emission_base_year = safe_float_convert(total_row['기준연도 배출량(tCO2eq)'])
        total_emission_previous_year = safe_float_convert(total_row['전년도 배출량(tCO2eq)'])
        
        # 비율 계산 (0으로 나누기 방지)
        total_emission_report_year_vs_base_year = round(
            (total_emission_report_year / total_emission_base_year * 100) if total_emission_base_year != 0 else 0, 2
        )
        total_emission_report_year_vs_previous_year = round(
            (total_emission_report_year / total_emission_previous_year * 100) if total_emission_previous_year != 0 else 0, 2
        )
        
        # Scope별 배출량 추출
        scope1_row = table1_df[(table1_df['구분'] == 'Scope 1') & (table1_df['세부구분'] == '합계')]
        scope2_row = table1_df[(table1_df['구분'] == 'Scope 2') & (table1_df['세부구분'] == '합계')]
        scope3_row = table1_df[(table1_df['구분'] == 'Scope 3') & (table1_df['세부구분'] == '합계')]
        
        scope1_emission_report_year = safe_float_convert(scope1_row.iloc[0]['보고대상연도 배출량(tCO2eq)']) if not scope1_row.empty else 0
        scope2_emission_report_year = safe_float_convert(scope2_row.iloc[0]['보고대상연도 배출량(tCO2eq)']) if not scope2_row.empty else 0
        scope3_emission_report_year = safe_float_convert(scope3_row.iloc[0]['보고대상연도 배출량(tCO2eq)']) if not scope3_row.empty else 0
        
        # Scope별 비율 계산
        scope1_emission_rate_report_year = round(
            (scope1_emission_report_year / total_emission_report_year * 100) if total_emission_report_year != 0 else 0, 2
        )
        scope2_emission_rate_report_year = round(
            (scope2_emission_report_year / total_emission_report_year * 100) if total_emission_report_year != 0 else 0, 2
        )
        scope3_emission_rate_report_year = round(
            (scope3_emission_report_year / total_emission_report_year * 100) if total_emission_report_year != 0 else 0, 2
        )
        
        # 최다 배출원 찾기 (Scope 1, 2에서)
        scope1_2_data = table1_df[
            (table1_df['구분'].isin(['Scope 1', 'Scope 2'])) & 
            (table1_df['세부구분'] != '합계')
        ].copy()
        
        if not scope1_2_data.empty:
            scope1_2_data['emission_value'] = scope1_2_data['보고대상연도 배출량(tCO2eq)'].apply(safe_float_convert)
            largest_source_row = scope1_2_data.loc[scope1_2_data['emission_value'].idxmax()]
            
            name_largest_emission_source_report_year = largest_source_row['세부구분']
            amount_largest_emission_source_report_year = largest_source_row['emission_value']
            rate_largest_emission_source_report_year = round(
                (amount_largest_emission_source_report_year / total_emission_report_year * 100) if total_emission_report_year != 0 else 0, 2
            )
        else:
            name_largest_emission_source_report_year = "N/A"
            amount_largest_emission_source_report_year = 0
            rate_largest_emission_source_report_year = 0
        
        # 최다 배출 사업장 (Table5에서 추출)
        print("Table5 데이터 파싱 중...")
        table5_df = csv_string_to_dataframe(request_data.word_table5_csv)
        print(f"Table5 DataFrame shape: {table5_df.shape}")
        
        if not table5_df.empty:
            total_row_table5 = table5_df[table5_df['구분'] == '총합계']
            if not total_row_table5.empty:
                total_row_table5 = total_row_table5.iloc[0]
                workplace_columns = [col for col in table5_df.columns if col not in ['구분', '세부구분', '합계']]
                workplace_emissions = {}
                
                for col in workplace_columns:
                    try:
                        emission_value = safe_float_convert(total_row_table5[col])
                        workplace_emissions[col] = emission_value
                    except:
                        workplace_emissions[col] = 0
                
                if workplace_emissions:
                    largest_workplace = max(workplace_emissions.items(), key=lambda x: x[1])
                    name_largest_emission_workplace_report_year = largest_workplace[0]
                    amount_largest_emission_workplace_report_year = largest_workplace[1]
                    rate_largest_emission_workplace_report_year = round(
                        (amount_largest_emission_workplace_report_year / total_emission_report_year * 100) if total_emission_report_year != 0 else 0, 2
                    )
                else:
                    name_largest_emission_workplace_report_year = "N/A"
                    amount_largest_emission_workplace_report_year = 0
                    rate_largest_emission_workplace_report_year = 0
            else:
                name_largest_emission_workplace_report_year = "N/A"
                amount_largest_emission_workplace_report_year = 0
                rate_largest_emission_workplace_report_year = 0
        else:
            name_largest_emission_workplace_report_year = "N/A"
            amount_largest_emission_workplace_report_year = 0
            rate_largest_emission_workplace_report_year = 0
        
        # 매출액 및 임직원수 관련 계산
        revenue_report_year = safe_float_convert(request_data.report_sales)
        num_employee_report_year = int(safe_float_convert(request_data.report_employees))
        revenue_previous_year = safe_float_convert(request_data.report_sales_last_year)
        num_employee_previous_year = int(safe_float_convert(request_data.report_employees_last_year))
        
        emission_vs_revenue_report_year = round(
            (total_emission_report_year / revenue_report_year) if revenue_report_year != 0 else 0, 4
        )
        emission_vs_employee_report_year = round(
            (total_emission_report_year / num_employee_report_year) if num_employee_report_year != 0 else 0, 2
        )
        emission_vs_revenue_previous_year = round(
            (total_emission_previous_year / revenue_previous_year) if revenue_previous_year != 0 else 0, 4
        )
        emission_vs_employee_previous_year = round(
            (total_emission_previous_year / num_employee_previous_year) if num_employee_previous_year != 0 else 0, 2
        )
        
        return {
            'company_name': company_name,
            'workplace_num': workplace_num,
            'report_year': report_year,
            'base_year': base_year,
            'previous_year': previous_year,
            'total_emission_report_year': total_emission_report_year,
            'total_emission_base_year': total_emission_base_year,
            'total_emission_previous_year': total_emission_previous_year,
            'total_emission_report_year_vs_base_year': total_emission_report_year_vs_base_year,
            'total_emission_report_year_vs_previous_year': total_emission_report_year_vs_previous_year,
            'scope1_emission_report_year': scope1_emission_report_year,
            'scope1_emission_rate_report_year': scope1_emission_rate_report_year,
            'scope2_emission_report_year': scope2_emission_report_year,
            'scope2_emission_rate_report_year': scope2_emission_rate_report_year,
            'scope3_emission_report_year': scope3_emission_report_year,
            'scope3_emission_rate_report_year': scope3_emission_rate_report_year,
            'name_largest_emission_source_report_year': name_largest_emission_source_report_year,
            'amount_largest_emission_source_report_year': amount_largest_emission_source_report_year,
            'rate_largest_emission_source_report_year': rate_largest_emission_source_report_year,
            'name_largest_emission_workplace_report_year': name_largest_emission_workplace_report_year,
            'amount_largest_emission_workplace_report_year': amount_largest_emission_workplace_report_year,
            'rate_largest_emission_workplace_report_year': rate_largest_emission_workplace_report_year,
            'revenue_report_year': revenue_report_year,
            'emission_vs_revenue_report_year': emission_vs_revenue_report_year,
            'num_employee_report_year': num_employee_report_year,
            'emission_vs_employee_report_year': emission_vs_employee_report_year,
            'revenue_previous_year': revenue_previous_year,
            'num_employee_previous_year': num_employee_previous_year,
            'emission_vs_revenue_previous_year': emission_vs_revenue_previous_year,
            'emission_vs_employee_previous_year': emission_vs_employee_previous_year
        }
    
    except Exception as e:
        print(f"calculate_variables에서 오류 발생: {str(e)}")
        import traceback
        traceback.print_exc()
        raise e

def add_table_from_dataframe(document, df: pd.DataFrame):
    """DataFrame을 워드 문서에 테이블로 추가합니다."""
    if df.empty:
        return
    
    # 테이블 생성 (헤더 포함)
    table = document.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = 'Table Grid'
    
    # 헤더 추가
    for j, column_name in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(column_name)
        # 셀 정렬 설정
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.bold = True
    
    # 데이터 추가
    for i, row in df.iterrows():
        for j, value in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = str(value) if pd.notna(value) else ""
            # 셀 정렬 설정
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    document.add_paragraph()  # 테이블 후 빈 줄 추가

def create_emission_chart(chart1_csv: str, temp_dir: str):
    """Chart1 CSV 데이터를 이용하여 누적 세로 막대형 차트를 생성합니다."""
    try:
        # Chart1 CSV 데이터 파싱
        df = csv_string_to_dataframe(chart1_csv)
        
        if df.empty:
            print("차트 데이터가 비어있습니다.")
            return None
        
        # 데이터 준비 - 각 Scope별로 연도별 배출량
        scopes = df.iloc[:, 0].tolist()  # Scope 1, Scope 2, Scope 3
        base_year_data = [safe_float_convert(x) for x in df.iloc[:, 1].tolist()]
        previous_year_data = [safe_float_convert(x) for x in df.iloc[:, 2].tolist()]
        report_year_data = [safe_float_convert(x) for x in df.iloc[:, 3].tolist()]
        
        # 한글 폰트 설정
        font_list = fm.findSystemFonts(fontpaths=None, fontext='ttf')
        korean_fonts = []
        for font in font_list:
            try:
                if any(korean_char in font for korean_char in ['AppleGothic', 'Nanum', 'Malgun', 'Dotum', 'Batang']):
                    korean_fonts.append(font)
            except:
                continue
        
        if korean_fonts:
            plt.rcParams['font.family'] = 'AppleGothic'
        else:
            plt.rcParams['font.family'] = 'DejaVu Sans'
        
        # 차트 생성
        fig, ax = plt.subplots(figsize=(10, 7))
        
        # x축 위치 설정 (연도별)
        x = [0, 1, 2]
        width = 0.3
        
        # 누적 막대 차트 생성
        # 기준연도
        ax.bar(x[0], base_year_data[0], width, label='Scope 1', color='#2E86AB', alpha=0.8)
        ax.bar(x[0], base_year_data[1], width, bottom=base_year_data[0], label='Scope 2', color='#A23B72', alpha=0.8)
        ax.bar(x[0], base_year_data[2], width, bottom=base_year_data[0]+base_year_data[1], label='Scope 3', color='#F18F01', alpha=0.8)
        
        # 전년도
        ax.bar(x[1], previous_year_data[0], width, label='_nolegend_', color='#2E86AB', alpha=0.8)
        ax.bar(x[1], previous_year_data[1], width, bottom=previous_year_data[0], label='_nolegend_', color='#A23B72', alpha=0.8)
        ax.bar(x[1], previous_year_data[2], width, bottom=previous_year_data[0]+previous_year_data[1], label='_nolegend_', color='#F18F01', alpha=0.8)
        
        # 보고대상연도
        ax.bar(x[2], report_year_data[0], width, label='_nolegend_', color='#2E86AB', alpha=0.8)
        ax.bar(x[2], report_year_data[1], width, bottom=report_year_data[0], label='_nolegend_', color='#A23B72', alpha=0.8)
        ax.bar(x[2], report_year_data[2], width, bottom=report_year_data[0]+report_year_data[1], label='_nolegend_', color='#F18F01', alpha=0.8)
        
        # 축 레이블 및 제목 설정
        ax.set_xlabel('연도', fontsize=12, fontweight='bold')
        ax.set_ylabel('배출량 (tCO2eq)', fontsize=12, fontweight='bold')
        ax.set_title('연도별 Scope 온실가스 배출량 추이', fontsize=14, fontweight='bold', pad=20)
        
        # x축 눈금 설정
        ax.set_xticks(x)
        ax.set_xticklabels(['기준연도', '전년도', '보고대상연도'])
        
        # 범례 설정
        ax.legend(loc='upper right', fontsize=10, bbox_to_anchor=(1.0, 1.0))
        
        # 그리드 추가
        ax.grid(True, axis='y', alpha=0.3)
        
        # y축 범위 설정
        max_total = max([
            sum(base_year_data),
            sum(previous_year_data),
            sum(report_year_data)
        ])
        if max_total > 0:
            ax.set_ylim(0, max_total * 1.1)
        
        # 막대 위에 값 표시
        year_totals = [sum(base_year_data), sum(previous_year_data), sum(report_year_data)]
        for i, total in enumerate(year_totals):
            if total > 0:
                ax.text(i, total + max_total * 0.02, f'{total:,.0f}', 
                       ha='center', va='bottom', fontsize=10, fontweight='bold', color='black')
        
        # 레이아웃 조정
        plt.tight_layout()
        plt.subplots_adjust(right=0.85)
        
        # 차트 저장
        chart_path = os.path.join(temp_dir, 'Chart1.png')
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        print("누적 세로 막대형 차트가 성공적으로 생성되었습니다.")
        return chart_path
        
    except Exception as e:
        print(f"차트 생성 중 오류가 발생했습니다: {str(e)}")
        return None

def create_emission_report(variables: Dict[str, Any], request_data: ReportRequest, temp_dir: str):
    """온실가스 배출량 보고서를 생성합니다."""
    doc = Document()
    
    # 페이지 여백 설정
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # 표지 페이지
    # CI 로고 추가 (기존 로고 파일이 있는 경우)
    if os.path.exists('ci_logo.png'):
        doc.add_picture('ci_logo.png', width=Inches(1.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 제목 추가
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_paragraph.add_run(f"\n\n\n\n\n\n{variables['company_name']}\n\n")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    
    subtitle_run = title_paragraph.add_run(f"{variables['report_year']}년 온실가스 배출량 보고서\n")
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.bold = True
    
    english_title_run = title_paragraph.add_run(f"{variables['report_year']} Greenhouse Gas Emission Report")
    english_title_run.font.size = Pt(18)
    english_title_run.font.bold = True
    
    # 페이지 나누기
    doc.add_page_break()
    
    # 본문 Main Section 1
    doc.add_heading('배출량 산정 개요', level=1)
    
    # Sub Section 1-1
    doc.add_heading('배출량 산정기준', level=2)
    criteria_text = f"""• 기업명: {variables['company_name']}
• 보고 대상기간: {variables['report_year']}.01.01. ~ {variables['report_year']}.12.31.
• 보고 대상범위: Scope 1, 2, 3
• 조직 경계: 본점 포함 총 {variables['workplace_num']}개 사업장
• 산정 기준:
  온실가스 배출량거래제의 배출량 보고 및 인증에 관한 지침 (환경부 고시 제2023-211호)
  The Greenhouse Gas Protocol. A Corporate Accounting and Reporting Standard
  Corporate Value Chain (Scope 3) Accounting and Reporting Standard
  Technical Guidance for Calculating Scope 3 Emissions (version 1.0)"""
    doc.add_paragraph(criteria_text)
    
    # Sub Section 1-2
    doc.add_heading('배출량 개요', level=2)
    overview_text = f"""•{variables['report_year']}년 Scope 1, 2, 3 배출량은 {variables['total_emission_report_year']:,.0f} tCO2eq 입니다. 이는 기준연도({variables['base_year']}년) 배출량 대비 {variables['total_emission_report_year_vs_base_year']}% 이며, 전년도({variables['previous_year']}년) 배출량 대비 {variables['total_emission_report_year_vs_previous_year']}% 입니다.
•{variables['report_year']}년 Scope 1 배출량은 {variables['scope1_emission_report_year']:,.0f} tCO2eq ({variables['scope1_emission_rate_report_year']}%), Scope 2 배출량은 {variables['scope2_emission_report_year']:,.0f} tCO2eq ({variables['scope2_emission_rate_report_year']}%), Scope 3 배출량은 {variables['scope3_emission_report_year']:,.0f} tCO2eq ({variables['scope3_emission_rate_report_year']}%)였습니다.
•{variables['report_year']}년 가장 많이 배출한 배출부문(활동)은 "{variables['name_largest_emission_source_report_year']}"으로 {variables['amount_largest_emission_source_report_year']:,.0f} tCO2eq를 배출하였으며 전체 배출량의 {variables['rate_largest_emission_source_report_year']}%를 차지했습니다. 가장 많이 배출한 사업장은 "{variables['name_largest_emission_workplace_report_year']}"으로 {variables['amount_largest_emission_workplace_report_year']:,.0f} tCO2eq를 배출하였으며 전체 배출량의 {variables['rate_largest_emission_workplace_report_year']}%를 차지했습니다.
•{variables['report_year']}년 매출액 대비 배출량은 {variables['emission_vs_revenue_report_year']} tCO2eq/백만원이며, 임직원수 대비 배출량은 {variables['emission_vs_employee_report_year']} tCO2eq/명입니다."""
    doc.add_paragraph(overview_text)
    
    # Sub Section 1-3
    doc.add_heading('전체 배출량', level=2)
    table1_df = csv_string_to_dataframe(request_data.word_table1_csv)
    add_table_from_dataframe(doc, table1_df)
    
    # Sub Section 1-4
    doc.add_heading('배출량 추이', level=2)
    chart_path = create_emission_chart(request_data.word_chart1_csv, temp_dir)
    if chart_path and os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 페이지 나누기
    doc.add_page_break()
    
    # 본문 Main Section 2
    doc.add_heading('배출량 산정 상세내역', level=1)
    
    # Sub Section 2-1
    doc.add_heading('1. 보고 대상 기간', level=2)
    period_text = f"보고 대상 기간은 {variables['report_year']}.01.01. ~ {variables['report_year']}.12.31. 입니다."
    doc.add_paragraph(period_text)
    
    # Sub Section 2-2
    doc.add_heading('2. 보고 대상 범위', level=2)
    scope_text = "보고 대상 범위는 Scope 1, Scope 2, Scope 3 입니다."
    doc.add_paragraph(scope_text)
    
    # Sub Section 2-3
    doc.add_heading('3. 보고 대상 온실가스', level=2)
    gas_text = """보고 대상 온실가스는 ⌜기후위기 대응을 위한 탄소중립∙녹색성장 기본법⌟상 6대 온실가스(CO2, CH4, N2O, HFCs, PFCs, SF6)입니다. 지구온난화지수(GWP)는 ⌜온실가스 배출권거래제의 배출량 보고 및 인증에 관한 지침⌟에 따라 IPCC 2차 보고서의 값(SAR)을 적용하였습니다."""
    doc.add_paragraph(gas_text)
    
    table2_df = csv_string_to_dataframe(request_data.word_table2_csv)
    add_table_from_dataframe(doc, table2_df)
    
    # Sub Section 2-4
    doc.add_heading('4. 조직 경계', level=2)
    boundary_text = f"{variables['company_name']} 내에서 총 {variables['workplace_num']}개 사업장을 대상으로 온실가스 배출량을 측정하였으며, 각 사업장의 조직 경계는 아래 표와 같습니다."
    doc.add_paragraph(boundary_text)
    
    table3_df = csv_string_to_dataframe(request_data.word_table3_csv)
    add_table_from_dataframe(doc, table3_df)
    
    # Sub Section 2-5
    doc.add_heading('5. 운영 경계', level=2)
    operation_text = f"운영 경계에 따라 {variables['company_name']}의 배출원은 직접 배출/흡수량(Scope 1), 에너지 간접 배출량(Scope 2) 및 그 밖의 간접 배출량(Scope 3)으로 분류되었으며, 세부적인 내용은 다음과 같습니다."
    doc.add_paragraph(operation_text)
    
    table4_df = csv_string_to_dataframe(request_data.word_table4_csv)
    add_table_from_dataframe(doc, table4_df)
    
    # Sub Section 2-6
    doc.add_heading('6. 사업장별 온실가스 배출량', level=2)
    workplace_text = "각 사업장별 온실가스 배출량은 다음과 같습니다.\n\n(단위: tCO2eq)"
    doc.add_paragraph(workplace_text)
    
    table5_df = csv_string_to_dataframe(request_data.word_table5_csv)
    add_table_from_dataframe(doc, table5_df)
    
    # Sub Section 2-7
    doc.add_heading('7. Scope별 온실가스 배출량', level=2)
    scope_detail_text = "각 Scope별 온실가스 배출량은 다음과 같습니다."
    doc.add_paragraph(scope_detail_text)
    
    table6_df = csv_string_to_dataframe(request_data.word_table6_csv)
    add_table_from_dataframe(doc, table6_df)
    
    return doc

@app.post("/generate-report")
async def generate_report(request: ReportRequest):
    """온실가스 배출량 보고서를 생성하고 파일을 반환합니다."""
    temp_dir = None
    try:
        # 임시 디렉토리 생성
        temp_dir = tempfile.mkdtemp()
        
        # 변수 계산
        variables = calculate_variables(request)
        
        # 보고서 생성
        doc = create_emission_report(variables, request, temp_dir)
        
        # 파일 저장
        filename = f"{variables['report_year']}년_온실가스_배출량_보고서.docx"
        file_path = os.path.join(temp_dir, filename)
        doc.save(file_path)
        
        # 파일 반환
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        # 에러 발생 시 임시 디렉토리 정리
        if temp_dir and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=f"보고서 생성 중 오류가 발생했습니다: {str(e)}")

@app.get("/")
async def root():
    """API 상태 확인"""
    return {"message": "온실가스 배출량 보고서 생성 API가 정상적으로 작동 중입니다."}

@app.get("/health")
async def health_check():
    """헬스 체크"""
    return {"status": "healthy"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)